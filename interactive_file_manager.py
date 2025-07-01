#!/usr/bin/env python3
"""
interactive_file_manager.py - Database-Integrated Interactive File Management
Version 2.0 - With Smart File Analysis and Enhanced UI
"""

import os
import sqlite3
import sys
from datetime import datetime
from pathlib import Path

# --- Configuration ---
# The main vault path, can be overridden by an environment variable
VAULT_PATH = os.environ.get("VAULT_PATH", "/media/amy/EXTERNAL/Vault_Master_Data")
DATABASE_PATH = os.path.join(VAULT_PATH, "_System/data/vault.db")

# --- Database Class ---
class VaultDatabase:
    """Handles all interactions with the SQLite database."""
    
    def __init__(self, db_path=DATABASE_PATH):
        self.db_path = db_path
        # Ensure the directory for the database exists
        os.makedirs(os.path.dirname(self.db_path), exist_ok=True)
        self.ensure_database_exists()
    
    def ensure_database_exists(self):
        """Ensures the database and required tables are created."""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS files (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        description TEXT,
                        tags TEXT,
                        status TEXT,
                        content_type TEXT,
                        sensitivity TEXT,
                        custom_fields TEXT, -- For notes and other metadata
                        date_created TEXT,
                        filename TEXT,
                        path TEXT UNIQUE -- Path should be unique
                    )
                """)
                conn.commit()
        except Exception as e:
            print(f"❌ Database setup error: {e}")
    
    def add_file_metadata(self, file_data):
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO files (description, tags, status, content_type, 
                                    sensitivity, custom_fields, date_created, filename, path)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, file_data)
                conn.commit()
                return cursor.lastrowid
        except sqlite3.IntegrityError:
            print(f"⚠️ Warning: A record for this file path already exists in the database.")
            return None
        except Exception as e:
            print(f"❌ Error adding file metadata: {e}")
            return None
    
    def get_file_metadata(self, filename=None, path=None):
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                if path:
                    cursor.execute("SELECT * FROM files WHERE path=?", (str(path),))
                elif filename:
                    cursor.execute("SELECT * FROM files WHERE filename=?", (filename,))
                else:
                    cursor.execute("SELECT * FROM files")
                return cursor.fetchall()
        except Exception as e:
            print(f"❌ Error retrieving metadata: {e}")
            return []
    
    def update_file_metadata(self, file_id, field, value):
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                # Use f-string carefully as field names are controlled internally
                cursor.execute(f"UPDATE files SET {field}=? WHERE id=?", (value, file_id))
                conn.commit()
                return True
        except Exception as e:
            print(f"❌ Error updating metadata: {e}")
            return False

# --- Main Application Class ---
class InteractiveFileManager:
    """Main interactive file management system with smart capabilities."""
    
    def __init__(self):
        self.db = VaultDatabase()
        self.vault_path = Path(VAULT_PATH)
    
    def display_banner(self):
        print("\n" + "="*60)
        print("🗂️  VAULT MASTER - Interactive File Manager v2.0")
        print("🚀  AI-Enhanced Workflow & Knowledge Management")
        print("="*60)
    
    # --- File Analysis Methods ---
    def analyze_file_content(self, file_path):
        file_path = Path(file_path)
        analysis = {
            'size_human': self.human_readable_size(file_path.stat().st_size),
            'modified': datetime.fromtimestamp(file_path.stat().st_mtime).strftime('%Y-%m-%d %H:%M'),
            'extension': file_path.suffix.lower(),
            'content_preview': 'N/A',
            'detected_title': '',
            'smart_category': 'unknown',
            'suggested_tags': []
        }
        try:
            if analysis['extension'] == '.pdf':
                analysis.update(self.analyze_pdf(file_path))
            elif analysis['extension'] in ['.docx', '.doc']:  # Add DOCX support
                analysis.update(self.analyze_docx(file_path))
            elif analysis['extension'] in ['.txt', '.md', '.py', '.js', '.json', '.sh', '.sql']:
                analysis.update(self.analyze_text_file(file_path))
        except Exception as e:
            print(f"⚠️ Could not fully analyze file content: {e}")

        # Smart category/tag detection from filename
        filename_lower = file_path.name.lower()
        if any(w in filename_lower for w in ['research', 'study', 'analysis']):
            analysis['smart_category'], analysis['suggested_tags'] = 'research_document', ['🔬 Research', '📊 Analysis']
        elif any(w in filename_lower for w in ['legal', 'contract', 'agreement']):
            analysis['smart_category'], analysis['suggested_tags'] = 'legal_document', ['⚖️ Legal', '📋 Official']
        elif any(w in filename_lower for w in ['business', 'plan', 'proposal']):
            analysis['smart_category'], analysis['suggested_tags'] = 'business_document', ['💼 Business', '🎯 Strategy']
        return analysis

    def human_readable_size(self, size_bytes):
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0: return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} PB"

    def analyze_pdf(self, file_path):
        result = {'content_preview': 'Could not read PDF.', 'detected_title': '', 'smart_category': 'document_pdf'}
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                if pdf_reader.metadata and pdf_reader.metadata.title:
                    result['detected_title'] = pdf_reader.metadata.title
                if len(pdf_reader.pages) > 0:
                    text = pdf_reader.pages[0].extract_text()
                    if text:
                        result['content_preview'] = text.strip()[:250] + "..."
        except ImportError:
            result['content_preview'] = "NOTE: Install PyPDF2 (`pip install PyPDF2`) for PDF content preview."
        except Exception as e:
            result['content_preview'] = f"PDF analysis failed: {e}"
        result['suggested_tags'] = ['📄 PDF', '📋 Document']
        return result

    def analyze_docx(self, file_path):
        """Extract content preview from DOCX files"""
        result = {
            'content_preview': 'Could not read DOCX.',
            'detected_title': '',
            'smart_category': 'document_docx'
        }
        try:
            import docx
            doc = docx.Document(file_path)
            # Extract title from first paragraph or filename
            if doc.paragraphs and doc.paragraphs[0].text.strip():
                first_para = doc.paragraphs[0].text.strip()
                if len(first_para) < 100:
                    result['detected_title'] = first_para
            # Extract content preview from first few paragraphs
            content_parts = []
            char_count = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                if text and char_count < 300:
                    content_parts.append(text)
                    char_count += len(text)
            if content_parts:
                result['content_preview'] = '\n'.join(content_parts)[:300] + "..."
        except ImportError:
            result['content_preview'] = "NOTE: Install python-docx (`pip install python-docx`) for DOCX content preview."
        except Exception as e:
            result['content_preview'] = f"DOCX analysis failed: {e}"
        result['suggested_tags'] = ['📄 Document', '💼 Business']
        return result

    def analyze_text_file(self, file_path):
        result = {'content_preview': '', 'detected_title': '', 'smart_category': 'text_file'}
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read(500)
                lines = content.split('\n')
                if lines and lines[0].strip():
                    first_line = lines[0].strip()
                    if first_line.startswith('#'): result['detected_title'] = first_line.replace('#', '').strip()
                    else: result['detected_title'] = first_line[:70]
                result['content_preview'] = content[:250].strip() + "..."
        except Exception as e:
            result['content_preview'] = f"Could not read file: {e}"
        return result
    
    # --- Interactive UI Methods ---
    def display_file_analysis(self, file_path, analysis):
        print("\n" + "🔍" + "="*59)
        print(f"  FILE ANALYSIS: {file_path.name}")
        print("="*60)
        print(f"  📏 Size: {analysis['size_human']}   |   📅 Modified: {analysis['modified']}")
        if analysis['detected_title']: print(f"  📋 Detected Title: {analysis['detected_title']}")
        if analysis['content_preview']:
            print("\n  📖 Content Preview:")
            print(  "  " + "-"*40)
            print(f"  {analysis['content_preview']}")
            print(  "  " + "-"*40)
        print("="*60)

    def interactive_tag_selection(self, suggested_tags=None):
        print("\n🏷️  Tag Selection System")
        presets = {'1': ['📄 Document', '💼 Business'], '2': ['🔬 Research', '📊 Analysis'], '3': ['⚖️ Legal', '📋 Official'], '4': ['💻 Code', '🔧 Technical'], '5': ['🎨 Creative', '💡 Ideas'], '6': ['📝 Notes', '🤔 Personal']}
        if suggested_tags: print(f"💡 AI Suggestions: {', '.join(suggested_tags)}")
        for k, v in presets.items(): print(f"  {k}. {', '.join(v)}")
        print("  9. Custom tags | 0. Skip")
        choice = input("👉 Choose preset, 9 for custom, or Enter to use AI suggestions: ").strip()
        if choice == '' and suggested_tags: return ', '.join(suggested_tags)
        if choice in presets: return ', '.join(presets[choice])
        if choice == '9': return input("🏷️  Enter custom tags (comma-separated): ").strip()
        return "Uncategorized"

    def smart_category_selection(self, detected_category):
        print("\n📂 Content Type (Category)")
        print(f"🤖 AI Detected: {detected_category}")
        return input(f"👉 Press Enter to accept, or enter a new category: ").strip() or detected_category
    
    # --- Main Workflow ---
    def process_new_file(self, file_path_str):
        file_path = Path(file_path_str).resolve()
        if not file_path.exists():
            print(f"❌ File not found: {file_path}"); return
        
        print(f"\n🆕 Processing new file: {file_path.name}")
        analysis = self.analyze_file_content(file_path)
        self.display_file_analysis(file_path, analysis)
        
        # Interactive prompts with smart defaults
        print("\n✍️  Please provide metadata for the database:")
        default_desc = analysis.get('detected_title') or file_path.name
        description = input(f"📝 Description (or Enter for '{default_desc}'): ").strip() or default_desc
        content_type = self.smart_category_selection(analysis['smart_category'])
        tags = self.interactive_tag_selection(analysis['suggested_tags'])
        
        print("\n🔐 Sensitivity Level: 1. Public, 2. Internal, 3. Confidential")
        sens_choice = input("👉 Choose (or Enter for 'Internal'): ").strip()
        sensitivity = {'1': 'public', '3': 'confidential'}.get(sens_choice, 'internal')
        
        print("\n📊 Status: 1. New, 2. In Progress, 3. Review, 4. Completed, 5. Archived")
        status_choice = input("👉 Choose (or Enter for 'New'): ").strip()
        status = {'2': 'in_progress', '3': 'review', '4': 'completed', '5':'archived'}.get(status_choice, 'new')

        custom_fields = input("🔧 Add any custom notes (optional): ").strip()

        print("\n" + "📋" + "="*59 + "\n  CONFIRM AND SAVE\n" + "="*60)
        print(f"  Desc: {description}\n  Type: {content_type}\n  Tags: {tags}\n  Sens: {sensitivity}\n  Stat: {status}")
        if input("\n💾 Save to database? (Y/n): ").strip().lower() in ['', 'y', 'yes']:
            file_data = (description, tags, status, content_type, sensitivity, custom_fields, datetime.now().isoformat(), file_path.name, str(file_path))
            if self.db.add_file_metadata(file_data):
                print("✅ File metadata saved successfully!")
            else:
                print("❌ Failed to save metadata.")
        else:
            print("❌ Operation cancelled.")
    
    def search_database(self, search_term):
        print(f"\n🔍 Searching for '{search_term}'...")
        results = self.db.get_file_metadata() # Simple search for now
        matches = [r for r in results if search_term.lower() in str(r).lower()]
        if not matches: print("❌ No results found."); return
        for r in matches: print(f"  ID {r[0]}: {r[8]} ({r[2]}) - {r[1]}")

    # --- Main Loop ---
    def main_menu(self):
        self.display_banner()
        while True:
            print("\n" + "🎯" + "="*59 + "\n  MAIN MENU\n" + "="*60)
            print("  1. Process New File  |  2. Search Database  |  3. Exit")
            choice = input("\n👉 Choose option: ").strip()
            if choice == '1':
                self.process_new_file(input("📁 Enter full file path: ").strip())
            elif choice == '2':
                self.search_database(input("🔍 Search for: ").strip())
            elif choice == '3':
                print("👋 Goodbye!"); break
            else:
                print("❌ Invalid choice.")

# --- Script Entry Point ---
def main():
    manager = InteractiveFileManager()
    # Non-interactive mode could be added here later
    manager.main_menu()

if __name__ == "__main__":
    main()
